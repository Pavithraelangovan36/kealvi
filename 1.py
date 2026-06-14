from docx import Document
from docx.shared import Pt

doc = Document()

title = doc.add_heading('Day 5 Assignment Report – Kealvi Polling Functionality', level=1)
p = doc.add_paragraph()
p.add_run('GitHub Repository: ').bold = True
p.add_run('https://github.com/Pavithraelangovan36/kealvi\n')
p.add_run('Vercel Deployment: ').bold = True
p.add_run('https://pavithra123-5qfy.vercel.app\n')
p.add_run('Supabase Project: ').bold = True
p.add_run('https://supabase.com/dashboard/project/lrpgockfpvxogjrrgvpv/editor/17558')

sections = [
    ("Introduction",
     "The objective of this assignment was to fork the Kealvi project, design and implement a polling system, create the required database schema in Supabase, connect the frontend with backend APIs, and deploy the application using Vercel."),
    ("Technologies Used",
     "Next.js, React.js, TypeScript, Supabase, PostgreSQL, GitHub, and Vercel."),
    ("Project Setup",
     "Forked the original repository, cloned it locally, installed dependencies, configured Supabase, created database tables, implemented API routes, integrated the frontend, and deployed the application on Vercel."),
    ("Database Design",
     "Polls Table: id (UUID), question (TEXT), created_at (TIMESTAMP). Poll Options Table: id (UUID), poll_id (UUID), option_text (TEXT), votes (INTEGER)."),
    ("Polling Features",
     "Create Poll, View Polls, Vote on Polls, and Display Poll Results dynamically."),
    ("API Endpoints",
     "POST /api/polls, GET /api/polls, POST /api/polls/vote, GET /api/polls/results."),
    ("Deployment",
     "The project was deployed successfully using Vercel for public access and testing."),
    ("Challenges Faced",
     "Database schema design, API integration, vote management, and frontend-backend connectivity."),
    ("Learning Outcomes",
     "Learned full-stack development, Supabase database management, REST API development, deployment, and GitHub workflow."),
    ("Conclusion",
     "The polling functionality was successfully implemented and integrated into the Kealvi project.")
]

for heading, text in sections:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(text)

path = "/mnt/data/Kealvi_Day5_Assignment_Report.docx"
doc.save(path)

print(path)
